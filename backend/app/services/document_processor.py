"""
Document processing service for multiple file formats
"""
import os
import hashlib
from pathlib import Path
from typing import Dict, Any, Optional
import logging

import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import markdown
from langchain.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process a document based on its type
        
        Args:
            file_path: Path to the document
            file_type: Type of document (pdf, docx, xlsx, txt, md)
            
        Returns:
            Dictionary containing processed content and metadata
        """
        try:
            processors = {
                'pdf': self._process_pdf,
                'docx': self._process_docx,
                'doc': self._process_docx,
                'xlsx': self._process_excel,
                'xls': self._process_excel,
                'txt': self._process_text,
                'md': self._process_markdown
            }
            
            processor = processors.get(file_type.lower())
            if not processor:
                raise ValueError(f"Unsupported file type: {file_type}")
                
            result = processor(file_path)
            
            # Add common metadata
            result['file_hash'] = self._calculate_file_hash(file_path)
            result['file_size'] = os.path.getsize(file_path)
            
            # Split content into chunks
            if result.get('content'):
                result['chunks'] = self.text_splitter.split_text(result['content'])
                result['chunk_count'] = len(result['chunks'])
                
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': None
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document"""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            # Extract text content
            content = "\n\n".join([doc.page_content for doc in documents])
            
            # Extract metadata
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'pdf_info': pdf_reader.metadata if pdf_reader.metadata else {}
                }
                
            return {
                'success': True,
                'content': content,
                'metadata': metadata,
                'documents': documents
            }
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract text with structure
            content_parts = []
            structure_info = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
                    
                    # Identify structure (headings)
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.split(' ')[-1] if ' ' in paragraph.style.name else '1'
                        structure_info.append({
                            'type': 'heading',
                            'level': level,
                            'text': paragraph.text
                        })
                    else:
                        structure_info.append({
                            'type': 'paragraph',
                            'text': paragraph.text[:100]  # First 100 chars
                        })
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_content.append(row_data)
                tables_data.append(table_content)
            
            content = "\n\n".join(content_parts)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'structure': structure_info[:10],  # First 10 structure elements
                'tables_summary': f"{len(tables_data)} tables found"
            }
            
            return {
                'success': True,
                'content': content,
                'metadata': metadata,
                'structure': structure_info,
                'tables': tables_data
            }
            
        except Exception as e:
            logger.error(f"Word document processing error: {str(e)}")
            raise
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            excel_file = pd.ExcelFile(file_path)
            sheets_content = []
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to descriptive text
                sheet_text = f"Worksheet: {sheet_name}\n"
                sheet_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                sheet_text += f"Columns: {', '.join(df.columns.tolist())}\n"
                
                if not df.empty:
                    sheet_text += "\nData Sample (first 5 rows):\n"
                    sheet_text += df.head(5).to_string(index=False)
                    
                sheets_content.append(sheet_text)
                sheets_data[sheet_name] = df.to_dict('records')
            
            content = "\n\n".join(sheets_content)
            
            metadata = {
                'sheet_count': len(excel_file.sheet_names),
                'sheets': excel_file.sheet_names,
                'total_rows': sum(len(pd.read_excel(file_path, sheet)) for sheet in excel_file.sheet_names)
            }
            
            return {
                'success': True,
                'content': content,
                'metadata': metadata,
                'sheets_data': sheets_data
            }
            
        except Exception as e:
            logger.error(f"Excel processing error: {str(e)}")
            raise
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            lines = content.split('\n')
            
            metadata = {
                'line_count': len(lines),
                'character_count': len(content),
                'word_count': len(content.split())
            }
            
            return {
                'success': True,
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Text file processing error: {str(e)}")
            raise
    
    def _process_markdown(self, file_path: str) -> Dict[str, Any]:
        """Process Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                
            # Convert to HTML for structure analysis
            html_content = markdown.markdown(md_content, extensions=['extra', 'toc'])
            
            # Extract headers
            import re
            headers = re.findall(r'^#{1,6}\s+(.+)$', md_content, re.MULTILINE)
            
            metadata = {
                'header_count': len(headers),
                'headers': headers[:10],  # First 10 headers
                'has_code_blocks': '```' in md_content,
                'has_tables': '|' in md_content
            }
            
            return {
                'success': True,
                'content': md_content,
                'html_content': html_content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Markdown processing error: {str(e)}")
            raise
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()